"""
Report Service - 报告生成服务
支持 Excel、Word、HTML、PDF、Markdown 等多种格式的测试报告导出
"""

import os
import json
import zipfile
from datetime import datetime
from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import send_file


class ReportService:
    """报告生成服务类"""
    
    def __init__(self, app=None):
        self.app = app
        self.export_dir = 'exports'
        os.makedirs(self.export_dir, exist_ok=True)
    
    def generate_excel_report(self, project_name: str, test_cases: list, executions: list, defects: list) -> str:
        """
        生成 Excel 格式测试报告
        
        Args:
            project_name: 项目名称
            test_cases: 测试用例列表
            executions: 执行记录列表
            defects: 缺陷列表
            
        Returns:
            生成的 Excel 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.xlsx'
        
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            # 测试用例 sheet
            if test_cases:
                tc_df = pd.DataFrame([{
                    'ID': tc.get('id', ''),
                    '标题': tc.get('title', ''),
                    '描述': tc.get('description', ''),
                    '优先级': tc.get('priority', ''),
                    '状态': tc.get('status', ''),
                    'AI 生成': '是' if tc.get('ai_generated') else '否',
                    '创建人': tc.get('created_by', ''),
                    '创建时间': tc.get('created_at', '')
                } for tc in test_cases])
                tc_df.to_excel(writer, sheet_name='测试用例', index=False)
            
            # 执行记录 sheet
            if executions:
                ex_df = pd.DataFrame([{
                    'ID': ex.get('id', ''),
                    '用例数': len(ex.get('test_case_ids', [])),
                    '执行人': ex.get('executor', ''),
                    '环境': ex.get('environment', ''),
                    '状态': ex.get('status', ''),
                    '进度': f"{ex.get('progress', 0)}%",
                    '开始时间': ex.get('started_at', ''),
                    '完成时间': ex.get('completed_at', ''),
                    '耗时 (秒)': ex.get('duration', 0)
                } for ex in executions])
                ex_df.to_excel(writer, sheet_name='执行记录', index=False)
            
            # 缺陷统计 sheet
            if defects:
                df_df = pd.DataFrame([{
                    'ID': df.get('id', ''),
                    '标题': df.get('title', ''),
                    '严重程度': df.get('severity', ''),
                    '状态': df.get('status', ''),
                    '指派给': df.get('assignee', ''),
                    'AI 分析': '是' if df.get('ai_analyzed') else '否',
                    '创建人': df.get('created_by', ''),
                    '创建时间': df.get('created_at', '')
                } for df in defects])
                df_df.to_excel(writer, sheet_name='缺陷列表', index=False)
                
                # 缺陷统计摘要
                severity_count = df_df['严重程度'].value_counts()
                status_count = df_df['状态'].value_counts()
                
                summary_data = {
                    '指标': ['严重缺陷', '高优缺陷', '中优缺陷', '低优缺陷', '新增', '处理中', '已解决', '已关闭'],
                    '数量': [
                        severity_count.get('critical', 0),
                        severity_count.get('high', 0),
                        severity_count.get('medium', 0),
                        severity_count.get('low', 0),
                        status_count.get('new', 0),
                        status_count.get('in_progress', 0),
                        status_count.get('resolved', 0),
                        status_count.get('closed', 0)
                    ]
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='缺陷统计', index=False)
        
        return filename
    
    def generate_pdf_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 PDF 格式测试报告（使用 WeasyPrint）
        
        Returns:
            生成的 PDF 文件路径
        """
        try:
            from weasyprint import HTML
            
            # 先生成 HTML 内容
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            html_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.html'
            pdf_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.pdf'
            
            # 生成 HTML 临时文件
            self.generate_html_report(project_name, test_cases, executions, defects, summary, output_path=html_filename)
            
            # 转换为 PDF
            html = HTML(filename=html_filename)
            html.write_pdf(pdf_filename)
            
            # 清理临时 HTML 文件
            os.remove(html_filename)
            
            return pdf_filename
            
        except ImportError:
            # WeasyPrint 未安装，降级为 HTML
            print('WeasyPrint 未安装，降级生成 HTML 报告')
            return self.generate_html_report(project_name, test_cases, executions, defects, summary)
    
    def generate_markdown_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 Markdown 格式测试报告
        
        Returns:
            生成的 Markdown 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.md'
        
        md_content = f'''# {project_name} 测试报告

**生成时间：** {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

## 📊 项目概要

| 指标 | 数值 |
|------|------|
| 总用例数 | {summary.get('total_cases', 0)} |
| 执行次数 | {summary.get('total_executions', 0)} |
| 缺陷总数 | {summary.get('total_defects', 0)} |
| 通过率 | {summary.get('pass_rate', 0):.1f}% |

---

## 📝 测试用例列表

| ID | 标题 | 优先级 | 状态 | AI 生成 |
|----|------|--------|------|--------|
'''
        
        for tc in test_cases[:50]:
            md_content += f"| {tc.get('id', '')} | {tc.get('title', '')} | {tc.get('priority', '')} | {tc.get('status', '')} | {'是' if tc.get('ai_generated') else '否'} |\n"
        
        md_content += '''
---

## 🐛 缺陷统计

### 按严重程度

'''
        
        severity_count = {}
        for df in defects:
            sev = df.get('severity', 'medium')
            severity_count[sev] = severity_count.get(sev, 0) + 1
        
        for sev, count in severity_count.items():
            md_content += f"- **{sev}**: {count}个\n"
        
        md_content += '''
### 按状态

'''
        
        status_count = {}
        for df in defects:
            sta = df.get('status', 'new')
            status_count[sta] = status_count.get(sta, 0) + 1
        
        for sta, count in status_count.items():
            md_content += f"- **{sta}**: {count}个\n"
        
        md_content += '''
---

## ▶️ 最近执行记录

| ID | 执行人 | 环境 | 状态 | 进度 | 耗时 |
|----|--------|------|------|------|------|
'''
        
        for ex in executions[-10:]:
            duration = ex.get('duration', 0)
            duration_str = f'{duration:.1f}s' if duration else '-'
            md_content += f"| {ex.get('id', '')} | {ex.get('executor', '')} | {ex.get('environment', '')} | {ex.get('status', '')} | {ex.get('progress', 0)}% | {duration_str} |\n"
        
        md_content += f'''
---

*本报告由 AI Test Hub 自动生成*
'''
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return filename
    
    def export_all_formats(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        一次性导出所有格式的报告，并打包为 ZIP
        
        Returns:
            ZIP 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}_all.zip'
        
        # 生成各种格式
        files_created = []
        
        excel_file = self.generate_excel_report(project_name, test_cases, executions, defects)
        files_created.append(excel_file)
        
        word_file = self.generate_word_report(project_name, test_cases, executions, defects, summary)
        files_created.append(word_file)
        
        html_file = self.generate_html_report(project_name, test_cases, executions, defects, summary)
        files_created.append(html_file)
        
        pdf_file = self.generate_pdf_report(project_name, test_cases, executions, defects, summary)
        files_created.append(pdf_file)
        
        md_file = self.generate_markdown_report(project_name, test_cases, executions, defects, summary)
        files_created.append(md_file)
        
        # 打包为 ZIP
        with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in files_created:
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname)
                # 清理临时文件
                os.remove(file_path)
        
        return zip_filename
    
    def generate_word_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 Word 格式测试报告
        
        Args:
            project_name: 项目名称
            test_cases: 测试用例列表
            executions: 执行记录列表
            defects: 缺陷列表
            summary: 项目汇总信息
            
        Returns:
            生成的 Word 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.docx'
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(f'{project_name} 测试报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 报告时间
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # 项目概要
        doc.add_heading('1. 项目概要', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        summary_data = [
            ('总用例数', str(summary.get('total_cases', 0))),
            ('执行次数', str(summary.get('total_executions', 0))),
            ('缺陷总数', str(summary.get('total_defects', 0))),
            ('通过率', f"{summary.get('pass_rate', 0):.1f}%")
        ]
        for i, (key, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = key
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # 测试用例概览
        doc.add_heading('2. 测试用例概览', level=1)
        if test_cases:
            tc_table = doc.add_table(rows=1, cols=5)
            tc_table.style = 'Table Grid'
            hdr_cells = tc_table.rows[0].cells
            headers = ['ID', '标题', '优先级', '状态', 'AI 生成']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            for tc in test_cases[:20]:  # 限制显示前 20 个
                row_cells = tc_table.add_row().cells
                row_cells[0].text = str(tc.get('id', ''))
                row_cells[1].text = tc.get('title', '')[:30] + '...' if len(tc.get('title', '')) > 30 else tc.get('title', '')
                row_cells[2].text = tc.get('priority', '')
                row_cells[3].text = tc.get('status', '')
                row_cells[4].text = '是' if tc.get('ai_generated') else '否'
        else:
            doc.add_paragraph('暂无测试用例')
        
        doc.add_paragraph()
        
        # 缺陷统计
        doc.add_heading('3. 缺陷统计', level=1)
        if defects:
            severity_count = {}
            status_count = {}
            for df in defects:
                sev = df.get('severity', 'medium')
                sta = df.get('status', 'new')
                severity_count[sev] = severity_count.get(sev, 0) + 1
                status_count[sta] = status_count.get(sta, 0) + 1
            
            doc.add_paragraph(f'**按严重程度统计：**')
            for sev, count in severity_count.items():
                doc.add_paragraph(f'  - {sev}: {count}个')
            
            doc.add_paragraph(f'**按状态统计：**')
            for sta, count in status_count.items():
                doc.add_paragraph(f'  - {sta}: {count}个')
        else:
            doc.add_paragraph('暂无缺陷记录')
        
        doc.add_paragraph()
        
        # 最近执行记录
        doc.add_heading('4. 最近执行记录', level=1)
        if executions:
            ex_table = doc.add_table(rows=1, cols=6)
            ex_table.style = 'Table Grid'
            hdr_cells = ex_table.rows[0].cells
            headers = ['ID', '执行人', '环境', '状态', '进度', '耗时']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            for ex in executions[-10:]:  # 最近 10 条
                row_cells = ex_table.add_row().cells
                row_cells[0].text = str(ex.get('id', ''))
                row_cells[1].text = ex.get('executor', '')
                row_cells[2].text = ex.get('environment', '')
                row_cells[3].text = ex.get('status', '')
                row_cells[4].text = f"{ex.get('progress', 0)}%"
                duration = ex.get('duration', 0)
                row_cells[5].text = f'{duration:.1f}s' if duration else '-'
        else:
            doc.add_paragraph('暂无执行记录')
        
        doc.save(filename)
        return filename
    
    def generate_pdf_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 PDF 格式测试报告（使用 WeasyPrint）
        
        Returns:
            生成的 PDF 文件路径
        """
        try:
            from weasyprint import HTML
            
            # 先生成 HTML 内容
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            html_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.html'
            pdf_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.pdf'
            
            # 生成 HTML 临时文件
            self.generate_html_report(project_name, test_cases, executions, defects, summary, output_path=html_filename)
            
            # 转换为 PDF
            html = HTML(filename=html_filename)
            html.write_pdf(pdf_filename)
            
            # 清理临时 HTML 文件
            os.remove(html_filename)
            
            return pdf_filename
            
        except ImportError:
            # WeasyPrint 未安装，降级为 HTML
            print('WeasyPrint 未安装，降级生成 HTML 报告')
            return self.generate_html_report(project_name, test_cases, executions, defects, summary)
    
    def generate_markdown_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 Markdown 格式测试报告
        
        Returns:
            生成的 Markdown 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.md'
        
        md_content = f'''# {project_name} 测试报告

**生成时间：** {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

## 📊 项目概要

| 指标 | 数值 |
|------|------|
| 总用例数 | {summary.get('total_cases', 0)} |
| 执行次数 | {summary.get('total_executions', 0)} |
| 缺陷总数 | {summary.get('total_defects', 0)} |
| 通过率 | {summary.get('pass_rate', 0):.1f}% |

---

## 📝 测试用例列表

| ID | 标题 | 优先级 | 状态 | AI 生成 |
|----|------|--------|------|--------|
'''
        
        for tc in test_cases[:50]:
            md_content += f"| {tc.get('id', '')} | {tc.get('title', '')} | {tc.get('priority', '')} | {tc.get('status', '')} | {'是' if tc.get('ai_generated') else '否'} |\n"
        
        md_content += '''
---

## 🐛 缺陷统计

### 按严重程度

'''
        
        severity_count = {}
        for df in defects:
            sev = df.get('severity', 'medium')
            severity_count[sev] = severity_count.get(sev, 0) + 1
        
        for sev, count in severity_count.items():
            md_content += f"- **{sev}**: {count}个\n"
        
        md_content += '''
### 按状态

'''
        
        status_count = {}
        for df in defects:
            sta = df.get('status', 'new')
            status_count[sta] = status_count.get(sta, 0) + 1
        
        for sta, count in status_count.items():
            md_content += f"- **{sta}**: {count}个\n"
        
        md_content += '''
---

## ▶️ 最近执行记录

| ID | 执行人 | 环境 | 状态 | 进度 | 耗时 |
|----|--------|------|------|------|------|
'''
        
        for ex in executions[-10:]:
            duration = ex.get('duration', 0)
            duration_str = f'{duration:.1f}s' if duration else '-'
            md_content += f"| {ex.get('id', '')} | {ex.get('executor', '')} | {ex.get('environment', '')} | {ex.get('status', '')} | {ex.get('progress', 0)}% | {duration_str} |\n"
        
        md_content += f'''
---

*本报告由 AI Test Hub 自动生成*
'''
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return filename
    
    def export_all_formats(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        一次性导出所有格式的报告，并打包为 ZIP
        
        Returns:
            ZIP 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}_all.zip'
        
        # 生成各种格式
        files_created = []
        
        excel_file = self.generate_excel_report(project_name, test_cases, executions, defects)
        files_created.append(excel_file)
        
        word_file = self.generate_word_report(project_name, test_cases, executions, defects, summary)
        files_created.append(word_file)
        
        html_file = self.generate_html_report(project_name, test_cases, executions, defects, summary)
        files_created.append(html_file)
        
        pdf_file = self.generate_pdf_report(project_name, test_cases, executions, defects, summary)
        files_created.append(pdf_file)
        
        md_file = self.generate_markdown_report(project_name, test_cases, executions, defects, summary)
        files_created.append(md_file)
        
        # 打包为 ZIP
        with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in files_created:
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname)
                # 清理临时文件
                os.remove(file_path)
        
        return zip_filename
    
    def generate_html_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 HTML 格式测试报告
        
        Returns:
            生成的 HTML 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.html'
        
        html_content = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project_name} 测试报告</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        @media print {{
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body class="bg-gray-50">
    <div class="max-w-7xl mx-auto p-6">
        <!-- 头部 -->
        <div class="bg-white rounded-lg shadow-md p-6 mb-6">
            <h1 class="text-3xl font-bold text-gray-800 mb-2">{project_name} 测试报告</h1>
            <p class="text-gray-600">生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        </div>
        
        <!-- 概要统计 -->
        <div class="grid grid-cols-1 md:grid-cols-4 gap-4 mb-6">
            <div class="bg-white rounded-lg shadow-md p-4">
                <div class="text-sm text-gray-600">总用例数</div>
                <div class="text-3xl font-bold text-blue-600">{summary.get('total_cases', 0)}</div>
            </div>
            <div class="bg-white rounded-lg shadow-md p-4">
                <div class="text-sm text-gray-600">执行次数</div>
                <div class="text-3xl font-bold text-green-600">{summary.get('total_executions', 0)}</div>
            </div>
            <div class="bg-white rounded-lg shadow-md p-4">
                <div class="text-sm text-gray-600">缺陷总数</div>
                <div class="text-3xl font-bold text-red-600">{summary.get('total_defects', 0)}</div>
            </div>
            <div class="bg-white rounded-lg shadow-md p-4">
                <div class="text-sm text-gray-600">通过率</div>
                <div class="text-3xl font-bold text-purple-600">{summary.get('pass_rate', 0):.1f}%</div>
            </div>
        </div>
        
        <!-- 测试用例 -->
        <div class="bg-white rounded-lg shadow-md p-6 mb-6">
            <h2 class="text-xl font-bold text-gray-800 mb-4">测试用例列表</h2>
            <div class="overflow-x-auto">
                <table class="min-w-full divide-y divide-gray-200">
                    <thead class="bg-gray-50">
                        <tr>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">ID</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">标题</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">优先级</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">状态</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">AI 生成</th>
                        </tr>
                    </thead>
                    <tbody class="bg-white divide-y divide-gray-200">
                        {''.join([f"""
                        <tr>
                            <td class="px-4 py-2 whitespace-nowrap text-sm text-gray-900">{tc.get('id', '')}</td>
                            <td class="px-4 py-2 text-sm text-gray-900">{tc.get('title', '')}</td>
                            <td class="px-4 py-2 whitespace-nowrap"><span class="px-2 inline-flex text-xs leading-5 font-semibold rounded-full {'bg-red-100 text-red-800' if tc.get('priority') == 'high' else 'bg-yellow-100 text-yellow-800' if tc.get('priority') == 'medium' else 'bg-blue-100 text-blue-800'}">{tc.get('priority', '')}</span></td>
                            <td class="px-4 py-2 whitespace-nowrap"><span class="px-2 inline-flex text-xs leading-5 font-semibold rounded-full {'bg-green-100 text-green-800' if tc.get('status') == 'active' else 'bg-gray-100 text-gray-800'}">{tc.get('status', '')}</span></td>
                            <td class="px-4 py-2 whitespace-nowrap text-sm text-gray-900">{'是' if tc.get('ai_generated') else '否'}</td>
                        </tr>
                        """ for tc in test_cases[:50]])}
                    </tbody>
                </table>
            </div>
        </div>
        
        <!-- 缺陷列表 -->
        <div class="bg-white rounded-lg shadow-md p-6 mb-6">
            <h2 class="text-xl font-bold text-gray-800 mb-4">缺陷列表</h2>
            <div class="overflow-x-auto">
                <table class="min-w-full divide-y divide-gray-200">
                    <thead class="bg-gray-50">
                        <tr>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">ID</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">标题</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">严重程度</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">状态</th>
                            <th class="px-4 py-2 text-left text-xs font-medium text-gray-500 uppercase">AI 分析</th>
                        </tr>
                    </thead>
                    <tbody class="bg-white divide-y divide-gray-200">
                        {''.join([f"""
                        <tr>
                            <td class="px-4 py-2 whitespace-nowrap text-sm text-gray-900">{df.get('id', '')}</td>
                            <td class="px-4 py-2 text-sm text-gray-900">{df.get('title', '')}</td>
                            <td class="px-4 py-2 whitespace-nowrap"><span class="px-2 inline-flex text-xs leading-5 font-semibold rounded-full {'bg-red-100 text-red-800' if df.get('severity') == 'critical' else 'bg-orange-100 text-orange-800' if df.get('severity') == 'high' else 'bg-yellow-100 text-yellow-800'}">{df.get('severity', '')}</span></td>
                            <td class="px-4 py-2 whitespace-nowrap"><span class="px-2 inline-flex text-xs leading-5 font-semibold rounded-full {'bg-blue-100 text-blue-800' if df.get('status') == 'new' else 'bg-yellow-100 text-yellow-800' if df.get('status') == 'in_progress' else 'bg-green-100 text-green-800'}">{df.get('status', '')}</span></td>
                            <td class="px-4 py-2 whitespace-nowrap text-sm text-gray-900">{'是' if df.get('ai_analyzed') else '否'}</td>
                        </tr>
                        """ for df in defects[:50]])}
                    </tbody>
                </table>
            </div>
        </div>
        
        <!-- 打印按钮 -->
        <div class="no-print text-center">
            <button onclick="window.print()" class="bg-blue-600 hover:bg-blue-700 text-white font-bold py-2 px-6 rounded-lg">
                打印报告
            </button>
        </div>
    </div>
</body>
</html>'''
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filename
    
    def generate_pdf_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 PDF 格式测试报告（使用 WeasyPrint）
        
        Returns:
            生成的 PDF 文件路径
        """
        try:
            from weasyprint import HTML
            
            # 先生成 HTML 内容
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            html_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.html'
            pdf_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.pdf'
            
            # 生成 HTML 临时文件
            self.generate_html_report(project_name, test_cases, executions, defects, summary, output_path=html_filename)
            
            # 转换为 PDF
            html = HTML(filename=html_filename)
            html.write_pdf(pdf_filename)
            
            # 清理临时 HTML 文件
            os.remove(html_filename)
            
            return pdf_filename
            
        except ImportError:
            # WeasyPrint 未安装，降级为 HTML
            print('WeasyPrint 未安装，降级生成 HTML 报告')
            return self.generate_html_report(project_name, test_cases, executions, defects, summary)
    
    def generate_markdown_report(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        生成 Markdown 格式测试报告
        
        Returns:
            生成的 Markdown 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}.md'
        
        md_content = f'''# {project_name} 测试报告

**生成时间：** {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

## 📊 项目概要

| 指标 | 数值 |
|------|------|
| 总用例数 | {summary.get('total_cases', 0)} |
| 执行次数 | {summary.get('total_executions', 0)} |
| 缺陷总数 | {summary.get('total_defects', 0)} |
| 通过率 | {summary.get('pass_rate', 0):.1f}% |

---

## 📝 测试用例列表

| ID | 标题 | 优先级 | 状态 | AI 生成 |
|----|------|--------|------|--------|
'''
        
        for tc in test_cases[:50]:
            md_content += f"| {tc.get('id', '')} | {tc.get('title', '')} | {tc.get('priority', '')} | {tc.get('status', '')} | {'是' if tc.get('ai_generated') else '否'} |\n"
        
        md_content += '''
---

## 🐛 缺陷统计

### 按严重程度

'''
        
        severity_count = {}
        for df in defects:
            sev = df.get('severity', 'medium')
            severity_count[sev] = severity_count.get(sev, 0) + 1
        
        for sev, count in severity_count.items():
            md_content += f"- **{sev}**: {count}个\n"
        
        md_content += '''
### 按状态

'''
        
        status_count = {}
        for df in defects:
            sta = df.get('status', 'new')
            status_count[sta] = status_count.get(sta, 0) + 1
        
        for sta, count in status_count.items():
            md_content += f"- **{sta}**: {count}个\n"
        
        md_content += '''
---

## ▶️ 最近执行记录

| ID | 执行人 | 环境 | 状态 | 进度 | 耗时 |
|----|--------|------|------|------|------|
'''
        
        for ex in executions[-10:]:
            duration = ex.get('duration', 0)
            duration_str = f'{duration:.1f}s' if duration else '-'
            md_content += f"| {ex.get('id', '')} | {ex.get('executor', '')} | {ex.get('environment', '')} | {ex.get('status', '')} | {ex.get('progress', 0)}% | {duration_str} |\n"
        
        md_content += f'''
---

*本报告由 AI Test Hub 自动生成*
'''
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return filename
    
    def export_all_formats(self, project_name: str, test_cases: list, executions: list, defects: list, summary: dict) -> str:
        """
        一次性导出所有格式的报告，并打包为 ZIP
        
        Returns:
            ZIP 文件路径
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f'{self.export_dir}/test_report_{project_name}_{timestamp}_all.zip'
        
        # 生成各种格式
        files_created = []
        
        excel_file = self.generate_excel_report(project_name, test_cases, executions, defects)
        files_created.append(excel_file)
        
        word_file = self.generate_word_report(project_name, test_cases, executions, defects, summary)
        files_created.append(word_file)
        
        html_file = self.generate_html_report(project_name, test_cases, executions, defects, summary)
        files_created.append(html_file)
        
        pdf_file = self.generate_pdf_report(project_name, test_cases, executions, defects, summary)
        files_created.append(pdf_file)
        
        md_file = self.generate_markdown_report(project_name, test_cases, executions, defects, summary)
        files_created.append(md_file)
        
        # 打包为 ZIP
        with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in files_created:
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname)
                # 清理临时文件
                os.remove(file_path)
        
        return zip_filename